#!/usr/bin/env python3
"""
Generate PCCC report by modifying the original .docx template.
Preserves all original formatting, fonts, styles — only replaces numeric data.
Uses paragraph-level text replacement since runs may split arbitrarily.
"""
import json
import re
import sys
from docx import Document

TEMPLATE_PATH = '/home/long/codes/web/pccc/_doc/long.docx'


def fmt_vn(n, decimals=2):
    """Format number to string with Vietnamese comma decimal separator."""
    if n is None:
        return '0'
    n = float(n)
    if abs(n - round(n)) < 1e-9:
        return str(int(round(n)))
    s = f'{n:.{decimals}f}'.rstrip('0').rstrip('.')
    return s.replace('.', ',')


def replace_para_text(paragraph, replacements):
    """
    Apply a list of (old_substring, new_substring) replacements to the paragraph.
    Rebuilds the runs so all formatting from the first run is preserved.
    """
    # Build full text
    full_text = ''.join(run.text for run in paragraph.runs)
    original = full_text

    for old, new in replacements:
        full_text = full_text.replace(old, new)

    if full_text == original:
        return

    # Clear all runs except the first, put all text into first run
    if paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ''
        paragraph.runs[0].text = full_text


def replace_para_regex(paragraph, pattern, replacement):
    """Apply regex replacement to the full paragraph text."""
    full_text = ''.join(run.text for run in paragraph.runs)
    new_text, count = re.subn(pattern, replacement, full_text)
    if count == 0 or new_text == full_text:
        return
    if paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ''
        paragraph.runs[0].text = new_text


def generate_report(data, output_path):
    p = data.get('params', {})
    r = data.get('results', {})

    def rv(d, key, default=0):
        v = d.get(key, default)
        return v.get('value', default) if isinstance(v, dict) else float(v)

    tbc = rv(p, 'tbc', 3)
    tcb = rv(p, 'tcb', 2)
    ttk = rv(p, 'ttk', 2)
    distL = rv(p, 'distL', 1.3)
    vxe = rv(p, 'vxe', 60)
    vl = rv(p, 'vl', 1.2)
    fdc = rv(p, 'fdc', 80)
    ict = rv(p, 'ict', 0.1)
    ql = rv(p, 'ql', 3.5)
    nlPerTruck = int(rv(p, 'nlPerTruck', 4))

    ttd = rv(r, 'ttd')
    ttdTotal = rv(r, 'ttdTotal')
    rl = rv(r, 'rl')
    fcc = rv(r, 'fcc', fdc)
    qct = rv(r, 'qct')
    nlCeil = int(rv(r, 'nl'))
    nTruck = int(rv(r, 'ntruck'))
    qlm = rv(r, 'qlm')
    nlmNozzle = int(rv(r, 'nlmNozzle'))
    nlmTruck = int(rv(r, 'nlmTruck'))
    totalNozzle = nlCeil + nlmNozzle

    S = {
        'tbc': fmt_vn(tbc), 'tbc2': f'{int(tbc):02d}',
        'tcb': fmt_vn(tcb), 'tcb2': f'{int(tcb):02d}',
        'ttk': fmt_vn(ttk),
        'L': fmt_vn(distL, 2),
        'vxe': fmt_vn(vxe, 0),
        'vl': fmt_vn(vl, 2),
        'fdc': fmt_vn(fdc, 0),
        'ict': fmt_vn(ict, 2),
        'ql': fmt_vn(ql, 2),
        'nl_truck': fmt_vn(nlPerTruck, 0),
        'nl_truck2': f'{nlPerTruck:02d}',
        'ttd': fmt_vn(ttd, 2),
        'ttdTotal': fmt_vn(ttdTotal, 2),
        'rl': fmt_vn(rl, 2),
        'fcc': fmt_vn(fcc, 0),
        'qct': fmt_vn(qct, 2),
        'nl': fmt_vn(nlCeil, 0), 'nl2': f'{nlCeil:02d}',
        'ntruck': fmt_vn(nTruck, 0), 'ntruck2': f'{nTruck:02d}',
        'qlm': fmt_vn(qlm, 2),
        'nlm': fmt_vn(nlmNozzle, 0), 'nlm2': f'{nlmNozzle:02d}',
        'nlm_truck': fmt_vn(nlmTruck, 0), 'nlm_truck2': f'{nlmTruck:02d}',
        'total': fmt_vn(totalNozzle, 0),
    }

    ttdCapped = min(ttdTotal, 10)
    S['ttd_cap'] = fmt_vn(ttdCapped, 2)
    S['nl_raw'] = fmt_vn(nlCeil / max(nlPerTruck, 1))
    S['nl_frac'] = fmt_vn(qct / max(ql, 0.1))
    S['nlm_raw'] = fmt_vn(qlm / max(nlPerTruck, 1))
    S['nlm_frac'] = fmt_vn(qlm / max(ql, 0.1))
    S['qlm_half'] = fmt_vn(qlm * 0.5, 2)

    doc = Document(TEMPLATE_PATH)

    # Each entry: (paragraph_index, [(old_str, new_str), ...])
    repl = {}

    def add(idx, old, new):
        repl.setdefault(idx, []).append((old, new))

    # ---------- Paragraph replacements ----------

    add(1,  '08h30', '08h30')  # keep time as-is

    # [3] Diện tích đám cháy khoảng 80m2;
    add(3,  '80m2', S['fdc'] + 'm2')

    # [8] Tbc: 03 phút.
    add(8,  '03 phút', S['tbc'] + ' phút')

    # [9] Tcb = 2 phút.
    add(9,  'Tcb = 2 phút', 'Tcb = ' + S['tcb'] + ' phút')

    # [11] Vxe = 60km/h
    add(11, 'Vxe = 60km/h', 'Vxe = ' + S['vxe'] + 'km/h')

    # [12] L = 1,3km
    add(12, 'L = 1,3km', 'L = ' + S['L'] + 'km')

    # [13] => Ttđ = L.60/Vxe = 1,3.60/60 = 1,3 phút.
    add(13, '1,3.60/60 = 1,3 phút',
         S['L'] + '.60/' + S['vxe'] + ' = ' + S['ttd'] + ' phút')

    # [14] Ttk = 2 phút
    add(14, 'Ttk = 2 phút', 'Ttk = ' + S['ttk'] + ' phút')

    # [15] Ttd = 3 + 2 + 1,3 + 2 = 8,3 phút
    add(15, 'Ttd = 3 + 2 + 1,3 + 2 = 8,3 phút',
         'Ttd = ' + S['tbc'] + ' + ' + S['tcb'] + ' + ' + S['ttd'] + ' + ' + S['ttk'] + ' = ' + S['ttdTotal'] + ' phút')

    # [16] ... được xác định là 8,3 phút.
    add(16, 'là 8,3 phút.', 'là ' + S['ttdTotal'] + ' phút.')

    # [18] Rl = 0,5.Ttd.Vl + Ttđ.Vl = 0,5.10.1,2 + 8,3.1,2 = 15,96m.
    add(18, '0,5.10.1,2 + 8,3.1,2 = 15,96m',
         '0,5.' + S['ttd_cap'] + '.' + S['vl'] + ' + ' + S['ttdTotal'] + '.' + S['vl'] + ' = ' + S['rl'] + 'm')

    # [21] Vl = 1,2 m/ph.
    add(21, 'Vl = 1,2 m/ph', 'Vl = ' + S['vl'] + ' m/ph')

    # [26] ... 8,3 phút ... 80m2.
    add(26, '8,3 phút', S['ttdTotal'] + ' phút')
    add(26, 'của 02 phòng 80m2', 'của 02 phòng ' + S['fdc'] + 'm2')

    # [28] Fcc = Fđc = 80m2.
    add(28, 'Fcc = Fđc = 80m2', 'Fcc = Fđc = ' + S['fcc'] + 'm2')

    # [49] Nccxe = NL/nl = 3/4 = 0,75 (làm tròn 01 xe)
    add(49, '3/4 = 0,75 (làm tròn 01 xe)',
         S['nl'] + '/' + S['nl_truck'] + ' = ' + S['nl_raw'] + ' (làm tròn ' + S['ntruck2'] + ' xe)')

    # [50] ... tối đa 04 lăng B.
    add(50, '04 lăng B.', S['nl_truck2'] + ' lăng B.')

    # [52] Nt = Nccxe = 1 tổ
    add(52, '= 1 tổ', '= ' + S['ntruck'] + ' tổ')

    # [55] Nlm = N/n = 2/4 = 0,5 (lấy tròn 01 xe)
    add(55, '2/4 = 0,5 (lấy tròn 01 xe)',
         S['qlm'] + '/' + S['nl_truck'] + ' = ' + S['nlm_raw'] + ' (lấy tròn ' + S['nlm_truck2'] + ' xe)')

    # [57] ... 04 lăng B, nl = 4
    add(57, '04 lăng B, nl = 4', S['nl_truck2'] + ' lăng B, nl = ' + S['nl_truck'])

    # [58] Nlm = Nlm    = 1 tổ
    add(58, '= 1 tổ', '= ' + S['nlm_truck'] + ' tổ')

    # [60] ... huy động 01 xe chữa cháy, 01 xe cứu thương, 01 xe bồn ...
    add(60, '01 xe chữa cháy', S['ntruck2'] + ' xe chữa cháy')

    # [63] Qct = Fcc . ict = 80.0,1 = 8 l/s
    add(63, '80.0,1 = 8 l/s', S['fcc'] + '.' + S['ict'] + ' = ' + S['qct'] + ' l/s')

    # [64] ict = 0,1 l/s.m2
    add(64, 'ict = 0,1 l/s', 'ict = ' + S['ict'] + ' l/s')

    # [65] NL1 = Qct/ql = 8/3,5 = 2,28 (lấy tròn 03 lăng B)
    add(65, '8/3,5 = 2,28 (lấy tròn 03 lăng B)',
         S['qct'] + '/' + S['ql'] + ' = ' + S['nl_frac'] + ' (lấy tròn ' + S['nl2'] + ' lăng B)')

    # [66] ql = 3,5 l/s
    add(66, 'ql = 3,5 l/s', 'ql = ' + S['ql'] + ' l/s')

    # [69] Qlm = 0,5 x Qct = 0,25 x 8 = 2/s
    add(69, '0,25 x 8 = 2/s', S['qlm_half'] + ' x ' + S['qct'] + ' = ' + S['qlm'] + '/s')

    # [70] ql = 3,5 l/s
    add(70, 'ql = 3,5 l/s', 'ql = ' + S['ql'] + ' l/s')

    # [71] Nlm = Qlm/nl = 2/3,5 = 0,6 (lấy tròn 01 lăng B)
    add(71, '2/3,5 = 0,6 (lấy tròn 01 lăng B)',
         S['qlm'] + '/' + S['ql'] + ' = ' + S['nlm_frac'] + ' (lấy tròn ' + S['nlm2'] + ' lăng B)')

    # [72] NlccB= 3 lăng B
    add(72, 'NlccB= 3 lăng B', 'NlccB= ' + S['nl'] + ' lăng B')

    # [73] Nllmb= 1 lăng B
    add(73, 'Nllmb= 1 lăng B', 'Nllmb= ' + S['nlm'] + ' lăng B')

    # [74] NlB= NlccB + Nllmb = 3 + 1 = 4 lăng B
    add(74, '3 + 1 = 4', S['nl'] + ' + ' + S['nlm'] + ' = ' + S['total'])

    # Apply replacements
    for idx, replacements in repl.items():
        if idx < len(doc.paragraphs):
            replace_para_text(doc.paragraphs[idx], replacements)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    raw = sys.stdin.read()
    data = json.loads(raw)
    output = data.get('outputPath', '/tmp/pccc_report.docx')
    result = generate_report(data, output)
    print(json.dumps({'success': True, 'path': result}))
